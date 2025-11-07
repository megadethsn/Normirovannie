import customtkinter as ctk
from datetime import datetime, timedelta
import locale
from tkinter import filedialog
import sys
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# Получаем абсолютный путь к директории с шаблонами
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')

def set_russian_locale():
    try:
        # Пробуем разные варианты для Windows
        locale_options = ['ru_RU', 'Russian', 'Russian_Russia.1251', 'rus']
        for loc in locale_options:
            try:
                locale.setlocale(locale.LC_TIME, loc)
                print(f"Установлена локаль: {loc}")
                return
            except locale.Error:
                continue
        print("Не удалось установить русскую локаль, используем системную")
    except Exception as e:
        print(f"Ошибка установки локали: {e}")

# Устанавливаем локаль
set_russian_locale()

# Настройка внешнего вида
ctk.set_appearance_mode("dark")  # Режим: "light", "dark", "system"
ctk.set_default_color_theme("blue")  # Темы: "blue", "green", "dark-blue"

class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Генератор нормированных заданий")
        self.geometry("1000x700")  
        
        # Главный контейнер для страниц
        self.container = ctk.CTkFrame(self)
        self.container.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Создание страниц
        self.pages = {
            "MainPage": MainPage(self.container, self),
            "Novoros": Novoros(self.container, self),
            'Astrakhan': Astrakhan(self.container, self),
            'Arhangelsk': Arhangelsk(self.container, self),
            'Bor': Bor(self.container, self),
            'Ekaterinburg': Ekaterinburg(self.container, self),
            'Habarovsk': Habarovsk(self.container, self),
            'Kaliningrad': Kaliningrad(self.container, self),
            'NN': NN(self.container, self),
            'Murmansk': Murmansk(self.container, self),
            'Nahodka': Nahodka(self.container, self),
            'PK': PK(self.container, self),
            'Rostov': Rostov(self.container, self),
            'Samara': Samara(self.container, self),
            'Sevastopol': Sevastopol(self.container, self),
            'Spb': Spb(self.container, self),
            'US': US(self.container, self),
            'Vladivostok': Vladivostok(self.container, self)
        }
        
        # Размещение всех страниц
        for page in self.pages.values():
            page.grid(row=0, column=0, sticky="nsew")
        
        # Конфигурация grid для контейнера
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)
        
        # Показываем главную страницу
        self.show_page("MainPage")
    
    def show_page(self, page_name):
        """Показать выбранную страницу"""
        page = self.pages.get(page_name)
        if page:
            page.tkraise()
            if hasattr(page, "on_show"):
                page.on_show()

class MultiSelectDropdown(ctk.CTkFrame):
    def __init__(self, master, values, **kwargs):
        super().__init__(master, **kwargs)
        self.values = values
        self.selected = []
        self.check_vars = {}
        
        self.toggle_btn = ctk.CTkButton(
            self,
            text="Выбрать сотрудников ▼",
            command=self.toggle_dropdown,
            anchor="w"
        )
        self.toggle_btn.pack(fill="x")
        
        self.dropdown_frame = ctk.CTkFrame(self)
        self.dropdown_shown = False
        
        for value in values:
            var = ctk.BooleanVar(value=False)
            self.check_vars[value] = var
            cb = ctk.CTkCheckBox(
                self.dropdown_frame,
                text=value,
                variable=var,
                command=lambda v=value: self.update_selection(v)
            )
            cb.pack(anchor="w", pady=2)

    def toggle_dropdown(self):
        if self.dropdown_shown:
            self.dropdown_frame.pack_forget()
            self.toggle_btn.configure(text="Выбрать сотрудников ▼")
        else:
            self.dropdown_frame.pack(fill="x", pady=5)
            self.toggle_btn.configure(text="Скрыть список ▲")
        self.dropdown_shown = not self.dropdown_shown

    def update_selection(self, value):
        if self.check_vars[value].get():
            if value not in self.selected:
                self.selected.append(value)
        else:
            if value in self.selected:
                self.selected.remove(value)
        self.update_button_text()

    def update_button_text(self):
        if len(self.selected) == 0:
            self.toggle_btn.configure(text="Выбрать сотрудников ▼")
        else:
            self.toggle_btn.configure(text=f"Выбрано: {len(self.selected)} ▼")

    def get_selected(self):
        if len(self.selected) == 1:
            return self.selected[0]
        elif len(self.selected) == 2:
            joined_string = '\n'.join(self.selected)
            return f"\n{joined_string}\n"
        else:
            return "\n".join(self.selected)

class BasePage(ctk.CTkFrame):
    """Базовый класс для всех страниц"""
    def __init__(self, parent, controller, worker_list=[], template='', name='', east=False, template_fizo_path = ''):
        super().__init__(parent)
        self.controller = controller
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1) 
        self.font = ('TimesNewRoman', 13)
        self.template = template
        self.template_fizo_path = template_fizo_path
        self.name = name
        self.east = east
        self.results = {}
        self.month_names = {
            1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
            5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
            9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
        }
        
        self.annotation_frame = ctk.CTkFrame(master=self)
        self.annotation_frame.grid(row=0, column=0, padx=20, pady=10, sticky="nsew")

        ctk.CTkLabel(self.annotation_frame, text='Выберите необходимые параметры', font=('TimesNewRoman', 20)).grid(row=1, column=0, columnspan=4, sticky='nsew')

        self.info_frame = ctk.CTkFrame(master=self)
        self.info_frame.grid(row=1, column=0, padx=20, pady=10, sticky='nsew')

        # Номер заявки
        ctk.CTkLabel(self.info_frame, text='Номер заявки:', font=self.font).grid(row=0, column=0, padx=5, pady=5, sticky='w')

        self.num_var = ctk.StringVar()
        self.num_entry = ctk.CTkEntry(self.info_frame, width=300, textvariable=self.num_var, placeholder_text='Номер заявки по нормированному заданию')
        self.num_entry.grid(row=0, column=1, columnspan=3, padx=5, pady=5, sticky='ew')

        # Дата формирования
        ctk.CTkLabel(self.info_frame, text='Дата формирования:', font=self.font).grid(row=1, column=0, padx=5, pady=5, sticky='w')
        self.ISSUE_DATE = ctk.StringVar()
        self.issue_date_entry = ctk.CTkEntry(self.info_frame, width=300, textvariable=self.ISSUE_DATE)
        now = datetime.now()
        formatted_date = f'«{now.day}» {self.month_names[now.month]} {now.year} г.'
        self.issue_date_entry.insert(0, formatted_date)
        self.issue_date_entry.grid(row=1, column=1, columnspan=3, padx=5, pady=5, sticky='ew')

        # Дата проверки
        ctk.CTkLabel(self.info_frame, text='Дата проведения проверки:', font=self.font).grid(row=2, column=0, padx=5, pady=5, sticky='w')

        self.WORK_DATE = ctk.StringVar()
        self.work_date_entry = ctk.CTkEntry(self.info_frame, width=300, textvariable=self.WORK_DATE)
        next_day = self.work_date(east=self.east)
        formatted_date_next = f'{next_day.day} {self.month_names[next_day.month]} {next_day.year} г.'
        self.work_date_entry.insert(0, formatted_date_next)
        self.work_date_entry.grid(row=2, column=1, columnspan=3, padx=5, pady=5, sticky='ew')

        # Время проверки
        ctk.CTkLabel(self.info_frame, text='Время проведения проверки:', font=self.font).grid(row=3, column=0, padx=5, pady=5, sticky='w')
        
        self.TIME_START = ctk.StringVar()
        self.entry_start_time = ctk.CTkEntry(self.info_frame, width=100, textvariable=self.TIME_START)
        self.entry_start_time.insert(0, '10:00')
        self.entry_start_time.grid(row=3, column=1, padx=5, pady=5, sticky='w')
        
        ctk.CTkLabel(self.info_frame, text='-').grid(row=3, column=2, padx=5, pady=5)
        
        self.TIME_END = ctk.StringVar()
        self.entry_end_time = ctk.CTkEntry(self.info_frame, width=100,textvariable=self.TIME_END)
        self.entry_end_time.insert(0, '16:00')
        self.entry_end_time.grid(row=3, column=3, padx=5, pady=5, sticky='w')

        #Типы проверок и УИНы
        self.pfo_var = ctk.BooleanVar(value=False)
        self.fizo_var = ctk.BooleanVar(value=False)
        self.zun_var = ctk.BooleanVar(value=False)

        self.pfo_value = ctk.StringVar()
        self.fizo_value = ctk.StringVar()
        self.zun_value = ctk.StringVar()

        self.cb_pfo = ctk.CTkCheckBox(self.info_frame, text='ПФО', variable=self.pfo_var, command=lambda: self.toggle_entry('pfo'))
        self.cb_pfo.grid(row=4, column=0, padx=5, pady=5, sticky='w')
        
        self.cb_fizo = ctk.CTkCheckBox(self.info_frame, text='ФИЗО', variable=self.fizo_var, command=lambda: self.toggle_entry('fizo'))
        self.cb_fizo.grid(row=5, column=0, padx=5, pady=5, sticky='w')
        

        self.cb_zun = ctk.CTkCheckBox(self.info_frame, text='ЗУН', variable=self.zun_var, command=lambda: self.toggle_entry('zun'))
        self.cb_zun.grid(row=6, column=0, padx=5, pady=5, sticky='w')

        self.entry_pfo = ctk.CTkEntry(self.info_frame, textvariable=self.pfo_value)
        self.entry_fizo = ctk.CTkEntry(self.info_frame, textvariable=self.fizo_value)
        self.entry_zun = ctk.CTkEntry(self.info_frame, textvariable=self.zun_value)

        # Поле для номера заявки ФИЗО (будет появляться/скрываться)
        self.fizo_number_label = ctk.CTkLabel(self.info_frame, text='Номер заявки для ФИЗО', font=self.font)
        self.fizo_number_var = ctk.StringVar()
        self.fizo_number_entry = ctk.CTkEntry(self.info_frame, textvariable=self.fizo_number_var)

        # Выбор сотрудников
        ctk.CTkLabel(self.info_frame, text='Выберите сотрудников:', font=self.font).grid(row=7, column=0, padx=5, pady=5, sticky='w')
        
        self.worker_list = worker_list
        self.workers_dropdown = MultiSelectDropdown(self.info_frame, self.worker_list)
        self.workers_dropdown.grid(row=7, column=1, columnspan=4, padx=5, pady=5, sticky='ew')

        #Выбор должностного лица
        ctk.CTkLabel(self.info_frame, text='Выберите должностное лицо:', font=self.font).grid(row=8, column=0, pady=5, sticky='w')

        self.chief_var = ctk.StringVar(value='')
        self.cb_bezzubcev = ctk.CTkCheckBox(self.info_frame, text='Беззубцев А.А.', font=self.font, variable=self.chief_var, onvalue='Беззубцев', offvalue='')
        self.cb_bezzubcev.grid(row=8, column=1, sticky='w')

        self.cb_aliabiev = ctk.CTkCheckBox(self.info_frame, text='Алябьев А.Б.', font=self.font, variable=self.chief_var, onvalue='Алябьев', offvalue='')
        self.cb_aliabiev.grid(row=8, column=2, sticky='w')

        self.cb_popirina = ctk.CTkCheckBox(self.info_frame, text='Попырина Е.М.', font=self.font, variable=self.chief_var, onvalue='Попырина', offvalue='')
        self.cb_popirina.grid(row=8, column=3, sticky='w')

        #Исполнитель
        ctk.CTkLabel(self.info_frame, text='Укажите исполнителя:', font=self.font).grid(row=9, column=0, sticky='w')
        self.issuer = ctk.StringVar(value='')

        self.cb_marusya = ctk.CTkCheckBox(self.info_frame, text='Разгуляева М.А.', font=self.font, variable=self.issuer, onvalue='Маруся', offvalue='')
        self.cb_marusya.grid(row=9, column=1, sticky='w')
        self.cb_vantuz = ctk.CTkCheckBox(self.info_frame, text='Воротилов И.И.', font=self.font, variable=self.issuer, onvalue='Вантуз', offvalue='')
        self.cb_vantuz.grid(row=9, column=2, sticky='w')
        self.cb_creator = ctk.CTkCheckBox(self.info_frame, text='Желудков А.В.', font=self.font, variable=self.issuer, onvalue='Создатель!!!', offvalue='')
        self.cb_creator.grid(row=9, column=3, sticky='w')
        

        # Настройка колонок
        for i in range(4):
            self.info_frame.grid_columnconfigure(i, weight=1 if i > 0 else 0)

        #Сохранение
        self.btn_save = ctk.CTkButton(self, text='Сформировать документ', font=self.font, command=self.save_data_and_form_doc)
        self.btn_save.grid(row=2, column=0, sticky='nsew')

        # Настройка навигации по Tab и стрелкам
        self.setup_navigation()

    def setup_navigation(self):
        """Настройка навигации - упрощенная версия"""
        print("🔧 Настройка навигации...")
        
        # Создаем список всех виджетов в правильном порядке
        self.navigation_widgets = [
            self.num_entry,
            self.issue_date_entry, 
            self.work_date_entry,
            self.entry_start_time,
            self.entry_end_time,
            self.cb_pfo, self.entry_pfo,
            self.cb_fizo, self.entry_fizo, 
            self.cb_zun, self.entry_zun,
            self.fizo_number_entry,
            self.workers_dropdown.toggle_btn,
            self.cb_bezzubcev, self.cb_aliabiev, self.cb_popirina,
            self.cb_marusya, self.cb_vantuz, self.cb_creator,
            self.btn_save
        ]
        
        print(f"📋 Всего виджетов для навигации: {len(self.navigation_widgets)}")
        
        # Привязываем события к каждому виджету
        for i, widget in enumerate(self.navigation_widgets):
            if widget:  # Проверяем что виджет существует
                widget.bind('<Up>', self.on_arrow_up)
                widget.bind('<Down>', self.on_arrow_down)
                print(f"   {i}: {type(widget).__name__} - привязано")
        
        # Отдельно привязываем буфер обмена к полям ввода
        entries = [w for w in self.navigation_widgets if hasattr(w, 'get')]
        for entry in entries:
            entry.bind('<Control-c>', self.on_copy)
            entry.bind('<Control-x>', self.on_cut)
            entry.bind('<Control-v>', self.on_paste)

    def on_arrow_up(self, event):
        """Обработка стрелки вверх"""
        print("⬆️ Стрелка вверх нажата")
        self.navigate(-1)
        return "break"

    def on_arrow_down(self, event):
        """Обработка стрелки вниз"""
        print("⬇️ Стрелка вниз нажата")
        self.navigate(1)
        return "break"

    def navigate(self, direction):
        """Навигация между виджетами"""
        try:
            current_focus = self.focus_get()
            print(f"🎯 Текущий фокус: {current_focus}")
            
            if current_focus in self.navigation_widgets:
                current_index = self.navigation_widgets.index(current_focus)
                new_index = (current_index + direction) % len(self.navigation_widgets)
                next_widget = self.navigation_widgets[new_index]
                
                print(f"🔀 Переход от {current_index} к {new_index}: {type(next_widget).__name__}")
                next_widget.focus_set()
                
            else:
                # Если фокус не на наших виджетах, фокусируемся на первом
                print("🔀 Фокус не на виджетах, переходим к первому")
                self.navigation_widgets[0].focus_set()
                
        except Exception as e:
            print(f"❌ Ошибка навигации: {e}")

    def on_copy(self, event):
        """Копирование"""
        try:
            widget = event.widget
            if widget.selection_present():
                selected_text = widget.selection_get()
                self.clipboard_clear()
                self.clipboard_append(selected_text)
                print("📋 Скопировано:", selected_text)
        except Exception as e:
            print(f"❌ Ошибка копирования: {e}")
        return "break"

    def on_cut(self, event):
        """Вырезание"""
        try:
            widget = event.widget
            if widget.selection_present():
                selected_text = widget.selection_get()
                self.clipboard_clear()
                self.clipboard_append(selected_text)
                widget.delete("sel.first", "sel.last")
                print("✂️ Вырезано:", selected_text)
        except Exception as e:
            print(f"❌ Ошибка вырезания: {e}")
        return "break"

    def on_paste(self, event):
        """Вставка"""
        try:
            widget = event.widget
            text = self.clipboard_get()
            widget.insert("insert", text)
            print("📎 Вставлено:", text)
        except Exception as e:
            print(f"❌ Ошибка вставки: {e}")
        return "break"

    def on_show(self):
        """Сброс полей при показе страницы"""
        self.reset_fields()

    def reset_fields(self):
        """Сброс всех полей ввода"""
        self.num_var.set('')
        self.ISSUE_DATE.set('')
        now = datetime.now()
        formatted_date = f'«{now.day}» {self.month_names[now.month]} {now.year} г.'
        self.issue_date_entry.delete(0, 'end')
        self.issue_date_entry.insert(0, formatted_date)
        
        self.WORK_DATE.set('')
        next_day = self.work_date(east=self.east)
        formatted_date_next = f'{next_day.day} {self.month_names[next_day.month]} {next_day.year} г.'
        self.work_date_entry.delete(0, 'end')
        self.work_date_entry.insert(0, formatted_date_next)
        
        self.TIME_START.set('10:00')
        self.TIME_END.set('16:00')
        
        self.pfo_var.set(False)
        self.fizo_var.set(False)
        self.zun_var.set(False)
        self.pfo_value.set('')
        self.fizo_value.set('')
        self.zun_value.set('')
        self.fizo_number_var.set('')
        
        self.entry_pfo.grid_forget()
        self.entry_fizo.grid_forget()
        self.entry_zun.grid_forget()
        self.fizo_number_label.grid_forget()
        self.fizo_number_entry.grid_forget()
        
        # Сброс выбора сотрудников
        for value in self.worker_list:
            if value in self.workers_dropdown.check_vars:
                self.workers_dropdown.check_vars[value].set(False)
        self.workers_dropdown.selected = []
        self.workers_dropdown.update_button_text()
        
        self.chief_var.set('')
        self.issuer.set('')

    #Функция отображения поля ввода по чекбоксам ПФО, ФИЗО, ЗУН
    def toggle_entry(self, checkbox):
        if checkbox == 'pfo':
            if self.pfo_var.get():
                self.entry_pfo.grid(row=4, column=1, columnspan=3, padx=5, pady=5, sticky='ew')
            else:
                self.entry_pfo.grid_forget()
        elif checkbox == 'fizo':
            if self.fizo_var.get():
                if self.template_fizo_path:
                    self.entry_fizo.grid(row=5, column=1, padx=5, pady=5, sticky='ew')
                    self.fizo_number_label.grid(row=5, column=2, sticky='ew')
                    self.fizo_number_entry.grid(row=5, column=3, sticky='ew')
                else:
                    self.entry_fizo.grid(row=5, column=1, columnspan=3, padx=5, pady=5, sticky='ew')
                    self.fizo_number_label.grid_forget()
                    self.fizo_number_entry.grid_forget()
            else:
                self.entry_fizo.grid_forget()
                self.fizo_number_label.grid_forget()
                self.fizo_number_entry.grid_forget()
        elif checkbox == 'zun':
            if self.zun_var.get():
                self.entry_zun.grid(row=6, column=1, columnspan=3, padx=5, pady=5, sticky='ew')
            else:
                self.entry_zun.grid_forget()
    
    #Экзамены без физры первая страница
    def get_exams(self):
        res = []

        if self.pfo_var.get():
            res.append('личностных (психофизиологических) качеств')
        if self.zun_var.get():
            res.append('знаний, умений и навыков')
        
        if len(res) == 1:
                return res[0]
        if len(res) == 2:
                return ', а также '.join(res)


    #Фукнция сбора УИНОВ
    def get_uins(self):
        total_uins = set()
        exams = {}
        type_of_exams = []
        if self.pfo_var.get():
            exams['личностных (психофизиологических) качеств'] = self.formate_uins(self.pfo_value.get())
            type_of_exams.append('личностных (психофизиологических) качеств')
        if self.fizo_var.get():
            exams['уровня физической подготовки'] = self.formate_uins(self.fizo_value.get())
            type_of_exams.append('уровня физической подготовки')
        if self.zun_var.get():
            exams['знаний, умений и навыков'] = self.formate_uins(self.zun_value.get())
            type_of_exams.append('знаний, умений и навыков')
        
        result_string = ''
        exams_string = ''
        for index, (key,value) in enumerate(exams.items(), 1):
            for item in value:
                total_uins.add(item)
            if len(exams) == 1:
                result_string += f'{key} аттестуемых лиц (УИН: {"; ".join(value)}).'
                break
            if index == len(exams):
                result_string += f'а также {key} аттестуемых лиц (УИН: {"; ".join(value)}).'
            else:
                result_string += f'{key} аттестуемых лиц (УИН: {"; ".join(value)}), '

        if len(type_of_exams) == 1:
                exams_string = type_of_exams[0]
        elif len(type_of_exams) == 2:
                exams_string =  f"{', а также '.join(type_of_exams)}"
        elif len(type_of_exams) == 3:
                exams_string = f'{type_of_exams[0]}, {type_of_exams[1]}, а также {type_of_exams[2]}'
        
        total_uins = sorted(list(total_uins))
        return total_uins, result_string, exams_string


        
    #Функция сбора и формирования уинов
    def formate_uins(self, string):
        if not string:
            return []
        cur_year = datetime.now().year - 2000
        start_uin = f'11{cur_year}20020'
        string = string.split(', ')
        result = []
        for item in string:
            if '-' in item:
                start, end = [int(num) for num in item.split('-')]
                for i in range(start, end + 1):
                    result.append(start_uin + str(i))
            if 'п' in item or 'g' in item:
                result.append(f'12{cur_year}20020' + item[1:])
            else:
                result.append(start_uin + item)
        return result
    
    #Чекбоксы подписанты
    def get_result_chief(self):
        selected = []

        if self.chief_var.get() == 'Беззубцев':
            selected.append('Начальник отдела проверок сил ОТБ')
            selected.append('А.А. Беззубцев')

        if self.chief_var.get() == 'Алябьев':
            selected.append('Заместитель начальника Службы')
            selected.append('А.Б. Алябьев')

        if self.chief_var.get() == 'Попырина':
            selected.append('Заместитель начальника отдела проверок сил ОТБ')
            selected.append('Е.М. Попырина')
        
        return selected
    
    #Функция чекбоксы исполнители
    def get_result_issuer(self):
        if self.issuer.get() == 'Вантуз':
            return 'Воротилов Иван Иванович', 'oa13@msecurity.ru'
        elif self.issuer.get() == 'Маруся':
            return 'Разгуляева Мария Андреевна', 'oa6@msecurity.ru'
        elif self.issuer.get() == 'Создатель!!!':
            return 'Желудков Андрей Викторович', 'oa15@msecurity.ru'
        return '', ''


    #Функция определения следующей даты
    def work_date(self, east=False):
        is_friday = datetime.now().isoweekday() == 5
        if is_friday and east:
            return datetime.now() + timedelta(days=4)
        elif east and datetime.now().isoweekday() == 4:
            return datetime.now() + timedelta(days=4)
        elif is_friday:
            return datetime.now() + timedelta(days=3)
        elif east:
            return datetime.now() + timedelta(days=2)
        return datetime.now() + timedelta(days=1)
    
    #Функция расчета затраченого времени
    def est_time(self):
        try:
            start = int(self.TIME_START.get().split(':')[0])
            end = int(self.TIME_END.get().split(':')[0])
            res = end - start
            if self.fizo_var.get() and self.template_fizo_path:
                res += 1
            return f'{res} часов' if res >= 5 else f'{res} часа'
        except:
            return '6 часов'
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """
        Конвертация DOCX в PDF
        """
        try:
            # Быстрая проверка файла
            if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
                return False

            import comtypes.client
            import time

            word = None
            doc = None

            try:
                # Создаем объект Word
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = False
                word.DisplayAlerts = False
                
                # Короткая пауза
                time.sleep(1)

                # Открываем документ
                doc_path = os.path.abspath(docx_path)
                doc = word.Documents.Open(doc_path)
                time.sleep(1)

                # Сохраняем как PDF
                pdf_path_abs = os.path.abspath(pdf_path)
                doc.SaveAs(pdf_path_abs, FileFormat=17)
                time.sleep(1)

                # Проверяем результат
                return os.path.exists(pdf_path_abs) and os.path.getsize(pdf_path_abs) > 0

            except Exception:
                return False

            finally:
                # Закрываем все
                try:
                    if doc:
                        doc.Close(SaveChanges=False)
                except:
                    pass
                try:
                    if word:
                        word.Quit()
                except:
                    pass

        except Exception:
            return False
    #Функция замены меток в документах
    def formate_docx(self, replacements_dict, template):
        edited_doc = Document(template)
    
        for table in edited_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{IDENTIFICATORS}}' in paragraph.text:
                            paragraph.paragraph_format.left_indent = Pt(0)
                            paragraph.paragraph_format.first_line_indent = Pt(0)
                        for key, value in replacements_dict.items():
                            if key in paragraph.text:
                                original_alignment = paragraph.alignment
                                
                                paragraph.text = paragraph.text.replace(key, value)
                                
                                paragraph.alignment = original_alignment
                                if key == '{{JOB_TITLE}}':
                                    table.autofit = True
                                
                                if key not in ('{{JOB_TITLE}}', '{{FULLNAME}}'):
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                for run in paragraph.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(13)
        
        # Обработка обычных параграфов
        for paragraph in edited_doc.paragraphs:
            for key, value in replacements_dict.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)   
                    if key in ['{{NUMBER}}', '{{DATE_OF_ISSUE}}', '{{NUMBER_FIZO}}']:
                        for run in paragraph.runs:
                            run.font.bold = True
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(13)
        
        # Обработка колонтитулов 
        for sect in edited_doc.sections:
            footer = sect.footer
            for paragraph in footer.paragraphs:
                for key, value in replacements_dict.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(9)
        
        return edited_doc
            
        

    def show_simple_popup(self, title, message):
        """Простое всплывающее окно с авто-закрытием через 2 секунды"""
        try:
            popup = ctk.CTkToplevel(self)
            popup.title(title)
            popup.geometry("400x120")
            popup.resizable(False, False)
            popup.transient(self)
            popup.grab_set()
            
            # Центрирование
            popup.update_idletasks()
            x = (popup.winfo_screenwidth() // 2) - (400 // 2)
            y = (popup.winfo_screenheight() // 2) - (120 // 2)
            popup.geometry(f"400x120+{x}+{y}")
            
            label = ctk.CTkLabel(popup, text=message, wraplength=380)
            label.pack(pady=30, padx=10)
            
            # Автоматическое закрытие через 2 секунды
            popup.after(2000, popup.destroy)
            
        except Exception as e:
            print(f"Ошибка создания popup: {e}")

    def save_data_and_form_doc(self):
        time = f'с {self.TIME_START.get()} час. до {self.TIME_END.get()} час.'
        total_uins, exams_str, types_str = self.get_uins()
        job_title, fullname = self.get_result_chief()
        issuer, email = self.get_result_issuer()
        ids = ('; '.join(total_uins) + '.').strip() if len(total_uins) > 1 else total_uins[0].strip()
        fizo_ids = '; '.join(self.formate_uins(self.fizo_value.get())) + '.' if len(self.formate_uins(self.fizo_value.get())) > 1 else self.formate_uins(self.fizo_value.get())
        
        # Получаем номер заявки ФИЗО, если есть
        fizo_number = self.fizo_number_var.get() if hasattr(self, 'fizo_number_var') and self.fizo_var.get() else ''

        self.results = {
            '{{NUMBER}}': self.num_var.get(),
            '{{NUMBER_FIZO}}': fizo_number,
            '{{DATE_OF_ISSUE}}': self.ISSUE_DATE.get().lower(),
            '{{DATE_TO_WORK}}': self.WORK_DATE.get().lower(),
            '{{TYPE_OF_EXAMS}}': types_str,
            '{{IDENTIFICATORS}}': ids,
            '{{EXAMS_P2}}': exams_str,
            '{{TIME}}': time,
            '{{EST_TIME}}': self.est_time(),
            '{{PEOPLE}}': self.workers_dropdown.get_selected(),
            '{{JOB_TITLE}}': job_title,
            '{{FULLNAME}}': fullname,
            '{{ISSUER}}': issuer,
            '{{EMAIL}}': email,
            '{{EXAMS}}': self.get_exams(),
            '{{FIZO_UIN}}':fizo_ids
        }
        
        try:
            # Проверяем существование основного шаблона
            if not os.path.exists(self.template):
                self.show_simple_popup("Ошибка", f"Основной шаблон не найден: {self.template}")
                return
            
            # Создаем основной документ
            edited_doc = self.formate_docx(self.results, self.template)
            default_filename = f'Нормированное задание на {self.work_date(self.east).strftime("%d.%m")} {self.name}'
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Сохранить документ как"
            )
            
            if not file_path:
                return
            
            # Сохраняем основной DOCX
            edited_doc.save(file_path)
            
            # Конвертируем основной в PDF
            pdf_file_path = file_path.replace('.docx', '.pdf')
            pdf_success = self.convert_to_pdf(file_path, pdf_file_path)
            
            # Показываем результат для основного документа
            if pdf_success:
                self.show_simple_popup("Успешно", "Документ и PDF сформированы")
            else:
                self.show_simple_popup("Успешно", "Документ сформирован")
            
            # Теперь создаем документ ФИЗО (если выбран чекбокс ФИЗО и есть шаблон)
            if self.fizo_var.get():
                print("🔔 Создание документа ФИЗО...")
                self.create_fizo_document()
            else:
                print("ℹ️ ФИЗО не выбран, пропускаем создание")
                
            # Финальное уведомление
            self.show_simple_popup("Готово", "Документы сформированы")
                
        except Exception as e:
            error_msg = f"❌ Ошибка при формировании документа: {e.args}"
            print(error_msg)
            self.show_simple_popup("Ошибка", "Ошибка при формировании документа")

    def create_fizo_document(self):
        """Создание документа ФИЗО с проверкой данных"""
        try:
            print("🔄 Проверка условий для ФИЗО...")
            print(f"   template_fizo_path: {self.template_fizo_path}")
            print(f"   fizo_var.get(): {self.fizo_var.get()}")
            
            if not self.template_fizo_path:
                print("❌ Путь к шаблону ФИЗО не указан")
                return
                
            if not os.path.exists(self.template_fizo_path):
                error_msg = f"❌ Шаблон ФИЗО не найден: {self.template_fizo_path}"
                print(error_msg)
                self.show_simple_popup("Ошибка", "Шаблон ФИЗО не найден")
                return
            
            print(f"✅ Шаблон ФИЗО найден: {self.template_fizo_path}")
            
            # ПРОВЕРЯЕМ И ЧИСТИМ ДАННЫЕ
            print("🔍 Проверка данных для ФИЗО...")
            clean_results = {}
            for key, value in self.results.items():
                if isinstance(value, list):
                    # Если значение - список, преобразуем в строку
                    clean_value = ', '.join(str(item) for item in value)
                    print(f"   ⚠️  Исправлен {key}: список -> строка")
                elif value is None:
                    clean_value = ''
                    print(f"   ⚠️  Исправлен {key}: None -> пустая строка")
                else:
                    clean_value = str(value)
                clean_results[key] = clean_value
                print(f"   {key}: {type(value)} -> {type(clean_value)}")
            
            # Создаем документ ФИЗО с очищенными данными
            fizo_doc = self.formate_docx(clean_results, self.template_fizo_path)
            fizo_default_filename = f'Заявка ФИЗО на {self.work_date(east=self.east).strftime("%d.%m")} {self.name}'
            
            print("💾 Запрос сохранения ФИЗО документа...")
            fizo_file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=fizo_default_filename,
                title="Сохранить заявку ФИЗО как"
            )
            
            print(f"📁 Полученный путь: {fizo_file_path}")
            
            if not fizo_file_path:
                print("❌ Пользователь отменил сохранение ФИЗО")
                return
            
            # Сохраняем ФИЗО DOCX
            fizo_doc.save(fizo_file_path)
            print(f"✅ Документ ФИЗО сохранен: {fizo_file_path}")
            
            # Конвертируем ФИЗО в PDF
            try:
                fizo_pdf_path = fizo_file_path.replace('.docx', '.pdf')
                print(f"🔄 Конвертация: {fizo_file_path} -> {fizo_pdf_path}")
                fizo_pdf_success = self.convert_to_pdf(fizo_file_path, fizo_pdf_path)
                
                if fizo_pdf_success:
                    print("✅ ФИЗО PDF создан успешно")
                else:
                    print("⚠️ ФИЗО PDF не создан, но DOCX сохранен")
            except Exception as pdf_error:
                print(f"⚠️ Ошибка конвертации в PDF: {pdf_error}")
                print("📄 DOCX файл сохранен, PDF не создан")
                    
        except Exception as e:
            error_msg = f"❌ Ошибка при создании ФИЗО: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            self.show_simple_popup("Ошибка ФИЗО", "Ошибка при создании заявки ФИЗО")
        
class MainPage(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # Заголовок
        self.label = ctk.CTkLabel(self, text="Главная страница", font=("TimesNewRoman", 24))
        self.label.grid(row=0, column=0, columnspan=3, pady=20, padx=20)
        
        # Кнопки для перехода на другие страницы
        buttons = [
            ("Новороссийск", "Novoros"),
            ('Астрахань', 'Astrakhan'),
            ('Архангельск', 'Arhangelsk'),
            ('Бор', 'Bor'),
            ('Екатеринбург', 'Ekaterinburg'),
            ('Хабаровск', 'Habarovsk'),
            ('Калининград', 'Kaliningrad'),
            ('Нижний Новгород', 'NN'),
            ('Мурманск', 'Murmansk'),
            ('Находка', 'Nahodka'),
            ('Петропавловск-Камчатский', 'PK'),
            ('Ростов-на-Дону', 'Rostov'),
            ('Самара', 'Samara'),
            ('Севастополь', 'Sevastopol'),
            ('Санкт-Петербург', 'Spb'),
            ('Южно-Сахалинск', 'US'),
            ('Владивосток', 'Vladivostok')
        ]
        
        
        for i, (text, page) in enumerate(buttons):
            row = i // 3 + 1  
            col = i % 3
            btn = ctk.CTkButton(
                self, 
                text=text,
                command=lambda p=page: controller.show_page(p),
                height=40,
                width=200
            )
            btn.grid(row=row, column=col, pady=10, padx=10, sticky='nsew')
        
        close_btn = ctk.CTkButton(
            self,
            text="Выход",
            command=self.controller.quit,
            fg_color="red",
            hover_color="darkred",
            height=40,
            width=200
        )
        close_btn.grid(row=len(buttons)//3 + 2, column=1, pady=20, padx=10, sticky='nsew')


class Novoros(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Novoros.docx')
        super().__init__(parent, controller, 
                        ['Лагутин Ростислав Борисович', 'Соловьев Сергей Викторович', 'Соловьев Артем Андреевич'], 
                        template_path, 'Новороссийск')

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Vladivostok(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Vladivostok.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Vladivostok_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Проценко Игорь Владимирович', 'Лабонин Илья Валентинович', 'Родионова Елена Ивановна'], 
                        template_path, 'Владивосток', east=True, template_fizo_path=template_fizo_path)

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class US(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'US.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'US_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Мельчарик Владимир Юрьевич', 'Шадрин Владимир Валерьевич'], 
                        template_path, 'Сахалин', east=True, template_fizo_path=template_fizo_path)

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Spb(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Spb.docx')
        super().__init__(parent, controller, 
                        ['Савельев Дмитрий Генрихович', 'Семенова Мария Сергеевна'], 
                        template_path, 'Спб')

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Sevastopol(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Sevastopol.docx')
        super().__init__(parent, controller, 
                        ['Кошелев Тимур Сергеевич', 'Кошелева Джеваире Айдеровна'], 
                        template_path, 'Севастополь')

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Samara(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Samara.docx')
        super().__init__(parent, controller, 
                        ['Башмакова Виктория Владимировна'], 
                        template_path, 'Самара')

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Rostov(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Rostov.docx')
        super().__init__(parent, controller, 
                        ['Евстратова Наталья Павловна', 'Язьков Анатолий Сергеевич'], 
                        template_path, 'Ростов')

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class PK(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'PK.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'PK_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Каушанов Александр Викторович', 'Еремеев Геннадий Гайсович', 'Зотов Станислав Викторович'], 
                        template_path, 'Петропавловск-Камчатский', east=True, template_fizo_path=template_fizo_path)

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Nahodka(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Nahodka.docx')
        super().__init__(parent, controller, 
                        ['Равнянский Константин Витальевич', 'Зубакин Эдуард Николаевич'], 
                        template_path, 'Находка', east=True)

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Murmansk(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Murmansk.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Murmansk_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Савельева Карина Дмитриевна', 'Пахомов Андрей Юрьевич'], 
                        template_path, 'Мурманск', template_fizo_path=template_fizo_path)

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class NN(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'NN.docx')
        super().__init__(parent, controller, 
                        ['Еркович Наталья Евгеньевна'], 
                        template_path, 'Нижний Новгород')

        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Astrakhan(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Astrakhan.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Astrakhan_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Шматова Раиса Анатольевна', 'Сивоконь Светлана Викторовна'], 
                        template_path, 'Астрахань', template_fizo_path=template_fizo_path)
        
        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Arhangelsk(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Arhangelsk.docx')
        super().__init__(parent, controller, 
                        ['Шубина Ксения Александровна'], 
                        template_path, 'Архангельск')
        
        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Bor(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Bor.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Bor_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Селюнин Александр Николаевич', 'Аладьин Алексей Николаевич'], 
                        template_path, 'Бор', template_fizo_path=template_fizo_path)
        
        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Ekaterinburg(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Ekaterinburg.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Ekaterinburg_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Рослякова Надежда Викторовна'], 
                        template_path, 'Екатеринбург', east=True, template_fizo_path=template_fizo_path)
        
        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Habarovsk(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Habarovsk.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Habarovsk_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Смирных Евгений Михайлович', 'Смирных Надежда Евгеньевна'], 
                        template_path, 'Хабаровск', east=True, template_fizo_path=template_fizo_path)
        
        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

class Kaliningrad(BasePage):
    def __init__(self, parent, controller):
        template_path = os.path.join(TEMPLATES_DIR, 'Kaliningrad.docx')
        template_fizo_path = os.path.join(TEMPLATES_DIR, 'Kaliningrad_FIZO.docx')
        super().__init__(parent, controller, 
                        ['Сироткин Сергей Николаевич'], 
                        template_path, 'Калининград', template_fizo_path=template_fizo_path)
        
        ctk.CTkButton(
            self, 
            text='Вернуться назад', 
            command=lambda: controller.show_page('MainPage')
        ).grid(row=3, column=0, pady=20, padx=20, sticky="ew")

if __name__ == "__main__":
    app = App()
    app.mainloop()
