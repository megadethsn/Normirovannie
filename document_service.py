import os
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_docx(replacements, template_path):
    edited_doc = Document(template_path)

    for table in edited_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{IDENTIFICATORS}}" in paragraph.text:
                        paragraph.paragraph_format.left_indent = Pt(0)
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            original_alignment = paragraph.alignment
                            paragraph.text = paragraph.text.replace(key, value)
                            paragraph.alignment = original_alignment
                            if key == "{{JOB_TITLE}}":
                                table.autofit = True
                            if key not in ("{{JOB_TITLE}}", "{{FULLNAME}}"):
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(13)

    for paragraph in edited_doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                if key in ["{{NUMBER}}", "{{DATE_OF_ISSUE}}", "{{NUMBER_FIZO}}"]:
                    for run in paragraph.runs:
                        run.font.bold = True
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(13)

    for section in edited_doc.sections:
        for paragraph in section.footer.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)

    return edited_doc


def convert_to_pdf(docx_path, pdf_path):
    try:
        if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
            return False

        import comtypes.client

        word = None
        doc = None

        try:
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            time.sleep(1)

            doc = word.Documents.Open(os.path.abspath(docx_path))
            time.sleep(1)
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            time.sleep(1)

            return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0
        except Exception:
            return False
        finally:
            try:
                if doc:
                    doc.Close(SaveChanges=False)
            except Exception:
                pass
            try:
                if word:
                    word.Quit()
            except Exception:
                pass
    except Exception:
        return False
